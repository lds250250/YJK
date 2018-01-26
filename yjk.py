#!/usr/bin/python
# -*- coding: utf-8 -*-
import re

from docx import Document
from docx.enum.text import WD_UNDERLINE
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from config import (BORUOCENG, BORUOCENGHAO, BORUOCENGHAO2, DIZHI, QIANGUDUAN,
                    QIANGUDUAN2, WDISPBI1PATTERN, WDISPBI2PATTERN,
                    WDISPBI3PATTERN, WDISPBI4PATTERN, WDISPPATTERN,
                    WMASSPATTERN, WMASSPATTERN2, WZQPATTERN, WZQPATTERN2,
                    ZHILIANGBI, DZ)


# pyinstaller --hidden-import=queue 1.py
class YJK(object):
    def __init__(self):
        self.AUTHOR = '李栋'
        self.BANBEN = 'Beta 1.0'
        self.JUDGEKUANGJIA = False  # 判断是否是框架结构，默认‘否’
        self.JUDGEQIANGUDUAN = False  # 判断是否有地下室，默认‘否’
        self.JUDGEBORUOCENG = False  # 判断是否有薄弱层，默认‘否’
        self.DIZHI = DIZHI
        self.WMASS_PATTERN = re.compile(WMASSPATTERN, re.S)
        self.WMASS_PATTERN2 = re.compile(WMASSPATTERN2, re.S)
        self.WZQ_PATTERN = re.compile(WZQPATTERN, re.S)
        self.WZQ_PATTERN2 = re.compile(WZQPATTERN2, re.S)
        self.WDISP_PATTERN = re.compile(WDISPPATTERN, re.S)
        self.WDISP_BI1_PATTERN = re.compile(WDISPBI1PATTERN, re.S)
        self.WDISP_BI2_PATTERN = re.compile(WDISPBI2PATTERN, re.S)
        self.WDISP_BI3_PATTERN = re.compile(WDISPBI3PATTERN, re.S)
        self.WDISP_BI4_PATTERN = re.compile(WDISPBI4PATTERN, re.S)

    def DocHeading(self):
        self.document = Document()
        self.document.add_heading(u'精简YJK自查计算书', 0)

    def Author(self):
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(f'Author:{self.AUTHOR}')
        run.font.size = Pt(10)
        run.font.name = u'楷体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')

    def TQWmass(self):
        wmassDZ = (f'{self.DIZHI}/设计结果/wmass.out')
        with open(wmassDZ, 'r') as f:
            data = f.read()
        return (data)

    def CLWmass1(self):
        wmass = []
        result = re.findall(self.WMASS_PATTERN, self.data_wmass)
        for i in result[0]:
            wmass.append(i.strip())
        return (wmass)

    def CLWmass2(self):
        wmass = []
        result = re.findall(self.WMASS_PATTERN2, self.data_wmass)
        for i in result[0]:
            wmass.append(i.strip())
        return (wmass)

    # 基本信息
    def RunJibenxinxi(self):
        paragraph = self.document.add_paragraph()
        self.document.add_heading(u'基本信息', 1)
        run = paragraph.add_run('本项目为')
        self.GeshiWenben(run)
        run = paragraph.add_run(self.wmass1[0])
        self.GeshiNeirong(run)
        run = paragraph.add_run(',主要材料是')
        self.GeshiWenben(run)
        run = paragraph.add_run(self.wmass1[1])
        self.GeshiNeirong(run)
        run = paragraph.add_run('，')
        self.GeshiWenben(run)
        if int(self.wmass1[2]) == 0:
            run = paragraph.add_run('无地下室。')
            self.GeshiNeirong(run)
        elif int(self.wmass1[2]) > 0:
            run = paragraph.add_run(self.wmass1[2])
            self.GeshiNeirong(run)
            run = paragraph.add_run('层地下室。')
            self.GeshiWenben(run)
        else:
            run = paragraph.add_run('地下室程序出错。')
            self.GeshiWenben(run)

    # 嵌固端设置
    def RunQianguduan(self):
        self.paragraph = self.document.add_paragraph()
        paragraph = self.paragraph
        QianGuDuan = None
        if int(self.wmass1[3]) == 0:
            QianGuDuan = '基础顶嵌固'
        elif int(self.wmass1[2]) == 0:
            QianGuDuan = '嵌固端有误！ERROR！'
        elif int(self.wmass1[2]) > 0 and int(self.wmass1[3]) > 0:
            if int(self.wmass1[2]) == int(self.wmass1[3]):
                QianGuDuan = '地下室顶板嵌固'
            elif int(self.wmass1[2]) > int(self.wmass1[3]):
                QianGuDuan = (
                    f'地下室{int(self.wmass1[2])-int(self.wmass1[3])}底板嵌固')
        else:
            QianGuDuan = '嵌固端程序出错'
        run = paragraph.add_run(QianGuDuan)
        self.GeshiNeirong(run)
        run = paragraph.add_run('。')
        self.GeshiWenben(run)

    # 裙房设置
    def RunQunfang(self):
        paragraph = self.paragraph
        QunFang = None
        if int(self.wmass1[4]) == 0:
            QunFang = '无裙房'
        elif int(self.wmass1[4]) > 0:
            QunFang = f'有{self.wmass1[4]}层裙房'
        else:
            QunFang = '裙房程序出错!'
        run = paragraph.add_run(QunFang)
        self.GeshiNeirong(run)
        run = paragraph.add_run('。')
        self.GeshiWenben(run)

    # P-Dlt设置
    def RunPDelt(self):
        paragraph = self.paragraph
        PDelt = None
        if self.wmass1[5] == '否':
            PDelt = '注意！未考虑P-Delt效应！'
        elif self.wmass1[5] == '是':
            PDelt = '已考虑P-Delt'
        else:
            PDelt = 'P-Delt裙房程序出错！'
        run = paragraph.add_run(PDelt)
        self.GeshiNeirong(run)
        run = paragraph.add_run('。')
        self.GeshiWenben(run)

    # 风荷载信息
    def RunFenghezai(self):
        self.document.add_heading(u'风荷载信息', 1)
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run('粗糙程度为')
        self.GeshiWenben(run)
        run = paragraph.add_run(self.wmass1[6])
        self.GeshiNeirong(run)
        run = paragraph.add_run('类，基本风压为')
        self.GeshiWenben(run)
        run = paragraph.add_run(self.wmass1[7])
        self.GeshiNeirong(run)
        run = paragraph.add_run('，\nX向基本周期为')
        self.GeshiWenben(run)
        run = paragraph.add_run(self.wmass1[8])
        self.GeshiNeirong(run)
        run = paragraph.add_run('，Y向基本周期为')
        self.GeshiWenben(run)
        run = paragraph.add_run(self.wmass1[9])
        self.GeshiNeirong(run)
        run = paragraph.add_run('（周期是否回代）')
        self.GeshiJinggao(run)
        run = paragraph.add_run('，\n风荷载效应放大系数为')
        self.GeshiWenben(run)
        run = paragraph.add_run({self.wmass1[10]})
        self.GeshiNeirong(run)
        run = paragraph.add_run('，')
        self.GeshiWenben(run)
        run = paragraph.add_run('（《高规》4.2.2条--当建筑高度大于60m时为1.1）')
        self.GeshiJinggao(run)

    # 地震信息1
    def RunDizheng1(self):
        self.document.add_heading(u'地震基本信息', 1)
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run('烈度为')
        self.GeshiWenben(run)
        run = paragraph.add_run(self.wmass1[12])
        self.GeshiNeirong(run)
        run = paragraph.add_run('，第')
        self.GeshiWenben(run)
        run = paragraph.add_run(self.wmass1[11])
        self.GeshiNeirong(run)
        run = paragraph.add_run('组，')
        self.GeshiWenben(run)
        run = paragraph.add_run(self.wmass1[13])
        self.GeshiNeirong(run)
        run = paragraph.add_run('类场地，特征周期为')
        self.GeshiWenben(run)
        run = paragraph.add_run(self.wmass1[14])
        self.GeshiNeirong(run)
        run = paragraph.add_run(f'\n（本项目为{self.wmass1[0]}）')
        self.GeshiJinggao(run)
        run = paragraph.add_run('周期折减系数为')
        self.GeshiWenben(run)
        run = paragraph.add_run(self.wmass1[15])
        self.GeshiNeirong(run)
        run = paragraph.add_run('，\n框架的抗震等级为')
        self.GeshiWenben(run)
        run = paragraph.add_run(self.wmass1[16])
        self.GeshiNeirong(run)
        run = paragraph.add_run('级，剪力墙的抗震等级为')
        self.GeshiWenben(run)
        run = paragraph.add_run(self.wmass1[17])
        self.GeshiNeirong(run)
        run = paragraph.add_run('级')
        self.GeshiWenben(run)

    # 偶然偏心
    def RunOuranpianxin(self):
        self.paragraph = self.document.add_paragraph()
        paragraph = self.paragraph
        if self.wmass1[18] == '是':
            run = paragraph.add_run('已考虑偶然偏心，')
            self.GeshiWenben(run)

        elif self.wmass1[18] == '否':
            run = paragraph.add_run('偶然偏心未考虑！')
            run.underline = WD_UNDERLINE.WAVY_HEAVY  # 波浪线
            self.GeshiJinggao(run)
        else:
            run = paragraph.add_run('偶然偏心程序出错！')
            self.GeshiERROR(run)

    # 双向地震
    def RunShuangxiangdizheng(self):
        paragraph = self.paragraph
        if self.wmass1[19] == '是':
            run = paragraph.add_run('已考虑双向地震，')
            self.GeshiWenben(run)
        elif self.wmass1[19] == '否':
            run = paragraph.add_run('双向地震未考虑！')
            run.underline = WD_UNDERLINE.WAVY_HEAVY  # 波浪线
            self.GeshiJinggao(run)
        else:
            run = paragraph.add_run('双向地震程序出错！')
            self.GeshiERROR(run)

    # 地震不利方向
    def RunBulifangxiang(self):
        paragraph = self.paragraph
        Zuibulifangxiang = None
        if self.wmass1[20] == '是':
            Zuibulifangxiang = '已考虑地震最不利方向，'
            run = paragraph.add_run(Zuibulifangxiang)
            self.GeshiWenben(run)
        elif self.wmass1[20] == '否':
            Zuibulifangxiang = '地震最不利方向未考虑！'
            run = paragraph.add_run(Zuibulifangxiang)
            run.underline = WD_UNDERLINE.WAVY_HEAVY  # 波浪线
            self.GeshiJinggao(run)
        else:
            Zuibulifangxiang = '地震最不利方向程序出错！'
            run = paragraph.add_run(Zuibulifangxiang)
            self.GeshiERROR(run)

    # 活荷载
    def RunHuohezai(self):
        self.document.add_heading(u'设计调整信息', 1)
        paragraph = self.document.add_paragraph()
        if int(self.wmass1[21]) == 0 and float(self.wmass1[22]) == 1:
            run = paragraph.add_run('未进行活荷载调整！')
            run.underline = WD_UNDERLINE.WAVY_HEAVY  # 波浪线
            self.GeshiJinggao(run)
        elif int(self.wmass1[21]) != 0 and float(self.wmass1[22]) > 1:
            run = paragraph.add_run('同时进行了‘活荷载不利布置’和‘活荷载放大’！')
            run.underline = WD_UNDERLINE.WAVY_HEAVY  # 波浪线
            self.GeshiJinggao(run)
        elif int(self.wmass1[21]) != 0:
            run = paragraph.add_run(self.wmass1[21])
            self.GeshiNeirong(run)
            run = paragraph.add_run('层进行了活荷载不利布置。')
            self.GeshiWenben(run)
        elif float(self.wmass1[22]) > 1:
            run = paragraph.add_run('活荷载放大系数为')
            self.GeshiWenben(run)
            run = paragraph.add_run(f'{self.wmass1[22]}。')
            self.GeshiNeirong(run)
        else:
            run = paragraph.add_run('地震最不利方向程序出错！')
            self.GeshiERROR(run)

    # 保护层
    def RunBaohuceng(self):
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run('梁保护层厚度为')
        self.GeshiWenben(run)
        run = paragraph.add_run(self.wmass1[23])
        self.GeshiNeirong(run)
        run = paragraph.add_run('mm，柱保护层厚度为')
        self.GeshiWenben(run)
        run = paragraph.add_run(self.wmass1[24])
        self.GeshiNeirong(run)
        run = paragraph.add_run('mm。')
        self.GeshiWenben(run)
        if int(self.wmass1[23]) < 25 or int(self.wmass1[24]) < 25:
            run = paragraph.add_run(
                '\n（《砼规》8.2.1-1条--保护层厚度不应小于钢筋直径，请确认钢筋直径小于25mm！）')
            self.GeshiJinggao(run)

    # 包络设计
    def RunBaoruosheji(self):
        paragraph = self.document.add_paragraph()
        if self.wmass1[25] == '是':
            run = paragraph.add_run('已考虑多塔包络，')
            self.GeshiNeirong(run)
        if self.wmass1[26] == '是':
            run = paragraph.add_run('已考虑框架-框剪包络，')
            self.GeshiNeirong(run)
        if self.wmass1[27] == '是':
            run = paragraph.add_run('已考虑多模型包络，')
            self.GeshiNeirong(run)
        if (self.wmass1[25] == '否' and self.wmass1[26] == '否'
                and self.wmass1[27]) == '否':
            run = paragraph.add_run('未考虑包络设计，')
            self.GeshiNeirong(run)

    # WZQ

    def TQWzq(self):
        wzqDZ = (f'{self.DIZHI}/设计结果/wzq.out')
        with open(wzqDZ, 'r') as f:
            data = f.read()
        return (data)

    def CLWzq(self, data_wzq):
        wzq = []
        result = re.findall(self.WZQ_PATTERN, data_wzq)
        for i in result[0]:
            wzq.append(i.strip())
        return (wzq)

    def RunNiuzhuan(self):
        self.document.add_heading(u'扭转信息', 1)
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(f'{self.wzq[0]}=')
        self.GeshiWenben(run)
        if float(self.wzq[1]) >= 0.9:
            run = paragraph.add_run(f'{self.wzq[1]}!周期比大于0.9！')
            self.GeshiCuowu(run)
        else:
            run = paragraph.add_run({self.wzq[1]})
            self.GeshiNeirong(run)

    def RunDizhengfangxiang(self):
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(f'{self.wzq[2]}=')
        self.GeshiWenben(run)
        run = paragraph.add_run(self.wzq[3])
        self.GeshiNeirong(run)
        if self.wmass2[0] == '否':
            run = paragraph.add_run('未考虑最不利地震方向！')
            self.GeshiJinggao(run)

    # WDISP

    def TQWdisp(self):
        wdispDZ = (f'{self.DIZHI}/设计结果/wdisp.out')
        with open(wdispDZ, 'r') as f:
            data = f.read()
        return (data)

    def CLWdisp(self, data_wdisp):
        wdispJIAO = []
        wdispBI = []
        result = re.findall(self.WDISP_PATTERN, data_wdisp)
        wdispBI.append(re.findall(self.WDISP_BI1_PATTERN, data_wdisp))
        wdispBI.append(re.findall(self.WDISP_BI2_PATTERN, data_wdisp))
        wdispBI.append(re.findall(self.WDISP_BI3_PATTERN, data_wdisp))
        wdispBI.append(re.findall(self.WDISP_BI4_PATTERN, data_wdisp))
        for i in result[0]:
            wdispJIAO.append(i.strip())
        return (wdispJIAO, wdispBI)

    # 位移角
    def RunWeiyijiao(self, wdispJIAO):
        self.document.add_page_break()
        self.document.add_heading(u'位移角&位移比', 1)
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(f'（本项目为{self.wmass1[0]}）')
        self.GeshiJinggao(run)
        table = self.document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '位移角'
        hdr_cells[1].text = 'X'
        hdr_cells[2].text = 'Y'
        # 地震
        row_cells = table.add_row().cells
        row_cells[0].text = '地震'
        row_cells[1].text = wdispJIAO[0]
        row_cells[2].text = wdispJIAO[1]
        # 风
        row_cells = table.add_row().cells
        row_cells[0].text = '风'
        row_cells[1].text = wdispJIAO[2]
        row_cells[2].text = wdispJIAO[3]

    # 位移比
    def RunWeiyibi(self, wdispBI):
        paragraph = self.document.add_paragraph()
        wdispBImax = []
        for item in wdispBI:
            wdispBImax.append(max(item))
        if float(max(wdispBImax)) < 1.2:
            run = paragraph.add_run('位移比均小于1.2')
            self.GeshiWenben(run)
        elif float(max(wdispBImax)) >= 1.5:
            run = paragraph.add_run('位移比大于1.5！！！')
            self.GeshiCuowu(run)
        elif float(max(wdispBImax)) >= 1.4:
            run = paragraph.add_run('位移比大于1.4！')
            run.underline = WD_UNDERLINE.WAVY_HEAVY  # 波浪线
            self.GeshiJinggao(run)
        elif float(max(wdispBImax)) >= 1.2:
            run = paragraph.add_run('位移比介于1.2~1.4之间')
            self.GeshiNeirong(run)
        else:
            run = paragraph.add_run('位移比程序出错！')
            self.GeshiERROR(run)

    # 重点信息
    def CLWmassZD(self):
        self.CLWmass_Kuangjia()
        self.CLWmass_Qianguduan()
        self.CLWmass_Boruoceng()
        self.CLWmass_Zhiliangbi()

    # 判断框架结构
    def CLWmass_Kuangjia(self):
        self.document.add_heading(u'重点问题', 1)
        if (self.wmass1[0]) == '框架结构':
            self.JUDGEKUANGJIA = True
        else:
            result = re.findall(re.compile(QIANGUDUAN2, re.S), self.data_wmass)
            if result[0] == '否':
                paragraph = self.document.add_paragraph()
                run = paragraph.add_run(
                    '底部嵌固楼层刚度比未执行《高规》3.5.2-2，\n请于‘计算参数’—>‘设计信息’—>‘薄弱层判断与调整’中修改'
                )
                self.GeshiCuowu(run)

    # 判断嵌固端刚度
    def CLWmass_Qianguduan(self):
        if int(self.wmass1[3]) > 0:
            self.JUDGEQIANGUDUAN = True
            self.RunQianguduangangdu()

    # 判断薄弱层
    def CLWmass_Boruoceng(self):
        boruoceng = []
        result = re.findall(re.compile(BORUOCENG, re.S), self.data_wmass)
        for i in result:
            boruoceng.append(i.strip())
        if float(max(boruoceng)) > 1.00:
            self.JUDGEBORUOCENG = True
            self.RunBoruoceng()

    # 判断质量比超限
    def CLWmass_Zhiliangbi(self):
        result = re.findall(re.compile(ZHILIANGBI), self.data_wmass)
        if result:
            self.RunZhiliangbi(result)

    # def RunCegang(self):
    #     pass

    # 嵌固端输出
    def RunQianguduangangdu(self):
        cengshu = int(self.wmass1[3]) + 1
        result = re.findall(re.compile(QIANGUDUAN, re.S), self.data_wmass)
        for item in result:
            qgd = []
            for i in item:
                qgd.append(i.strip())
            if int(qgd[0]) == cengshu:
                paragraph = self.document.add_paragraph()
                run = paragraph.add_run(qgd[1])
                self.GeshiNeirong(run)
                run = paragraph.add_run('塔嵌固端刚度比X方向为：')
                self.GeshiWenben(run)
                run = paragraph.add_run(qgd[2])
                self.GeshiNeirong(run)
                run = paragraph.add_run('、Y方向为：')
                self.GeshiWenben(run)
                run = paragraph.add_run(qgd[3])
                self.GeshiNeirong(run)
                run = paragraph.add_run('。')
                self.GeshiWenben(run)
                if float(qgd[2]) >= 0.5 or float(qgd[3]) >= 0.5:
                    run = paragraph.add_run('\n刚度比大于0.5！不能作为嵌固端！')
                    self.GeshiCuowu(run)
                else:
                    run = paragraph.add_run('刚度比满足嵌固要求！')
                    self.GeshiNeirong(run)

    # 薄弱层输出
    def RunBoruoceng(self):
        if self.JUDGEBORUOCENG:
            self.RunBoruocenghao()
        else:
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run('本项目无薄弱层！')
            self.GeshiNeirong(run)

    # 输出薄弱层层号
    def RunBoruocenghao(self):
        txt = re.findall(re.compile(BORUOCENGHAO, re.S), self.data_wmass)
        result = re.findall(re.compile(BORUOCENGHAO2, re.S), txt[0])
        for item in result:
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run(f'{item[1]}塔第{item[0]}层为薄弱层！')
            self.GeshiCuowu(run)

    # 输出质量比超限层号
    def RunZhiliangbi(self, result):
        for item in result:
            paragraph = self.document.add_paragraph()
            run = paragraph.add_run(
                f'{item[1]}塔第{item[0]}层质量比>1.5 不满足《高规》3.5.6!')
            self.GeshiCuowu(run)

    # 标题格式
    def GeshiBiaoti(self, run):
        run.font.name = u'楷体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')

    # 文本格式
    def GeshiWenben(self, run):
        run.font.name = u'宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 内容格式
    def GeshiNeirong(self, run):
        run.bold = True  # 加粗
        run.underline = WD_UNDERLINE.WAVY_HEAVY  # 波浪线
        run.font.name = u'黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

    # 警告格式
    def GeshiJinggao(self, run):
        run.bold = True  # 加粗
        run.font.color.rgb = RGBColor(0xE3, 0x6C, 0x0A)  # 颜色
        run.font.name = u'黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

    # 错误内容
    def GeshiCuowu(self, run):
        run.bold = True  # 加粗
        run.underline = WD_UNDERLINE.WAVY_HEAVY  # 波浪线
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 颜色
        run.font.name = u'黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

    # 错误格式
    def GeshiERROR(self, run):
        run.bold = True  # 加粗
        run.underline = WD_UNDERLINE.THICK  # 下划线
        run.font.name = u'黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

    # wmass_go
    def goWmass(self):
        self.data_wmass = self.TQWmass()
        self.wmass1 = self.CLWmass1()
        self.wmass2 = self.CLWmass2()
        self.RunJibenxinxi()
        self.RunQianguduan()
        self.RunQunfang()
        self.RunPDelt()
        self.RunFenghezai()
        self.RunDizheng1()
        self.RunOuranpianxin()
        self.RunShuangxiangdizheng()
        self.RunBulifangxiang()
        self.RunHuohezai()
        self.RunBaohuceng()
        self.RunBaoruosheji()

    def gowzq(self):
        data_wzq = self.TQWzq()
        self.wzq = self.CLWzq(data_wzq)
        self.RunNiuzhuan()
        self.RunDizhengfangxiang()

    def gowdisp(self):
        data_wdisp = self.TQWdisp()
        wdispJIAO, wdispBI = self.CLWdisp(data_wdisp)
        self.RunWeiyijiao(wdispJIAO)
        self.RunWeiyibi(wdispBI)

    def gozhongdian(self):
        self.CLWmassZD()

    def AuthorShengming(self):
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(f'Author:{self.AUTHOR}')
        run.font.size = Pt(10)
        run.font.name = u'楷体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(f'版本:{self.BANBEN}')
        run.font.size = Pt(10)
        run.font.name = u'楷体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'楷体')

    def go(self):
        self.DocHeading()
        # self.Author()
        self.goWmass()
        self.gowzq()
        self.gowdisp()
        self.gozhongdian()
        self.AuthorShengming()
        self.document.save(f'{DZ}YJK自查计算书.docx')


if __name__ == '__main__':
    yjk = YJK()
    yjk.go()

# pyinstaller --hidden-import=queue 1.py
